"""Generate BGM from EP30 storyboard (.docx) with progress output."""
import sys
sys.path.insert(0, r'd:\1 a universe\BGM-creater')
sys.stdout.reconfigure(encoding='utf-8')

from pathlib import Path
from docx import Document
from bgm_v2 import generate_bgm_multisegment
from bgm_v2.config import Config
from bgm_v2.storyboard import analyze_storyboard

def read_docx_storyboard(file_path: str) -> str:
    """Read storyboard from .docx table format."""
    doc = Document(file_path)
    if not doc.tables:
        raise ValueError("No tables found in docx")

    table = doc.tables[0]
    rows = []

    for row in table.rows[1:]:  # 跳过表头行
        cells = [cell.text.strip() for cell in row.cells]
        if len(cells) >= 7 and cells[0].isdigit():  # 确保是数据行
            shot_num, desc,台词, camera, visual, notes, duration = cells
            rows.append(f'第{cells[0]}场 | {duration}s | {desc}')

    return '\n'.join(rows)

def main():
    docx_path = r"C:\Users\99384\Desktop\第 30 集分镜.docx"
    output_dir = Path(r"d:\1 a universe\BGM-creater\output\ep30_docx")
    output_dir.mkdir(parents=True, exist_ok=True)

    print("=" * 80, flush=True)
    print("EP30 分镜 BGM 生成", flush=True)
    print("=" * 80, flush=True)

    # 读取docx
    print(f"\n[1/3] 读取分镜文件...", flush=True)
    storyboard_text = read_docx_storyboard(docx_path)
    print(f"✅ 读取完成 ({len(storyboard_text)} 字符)", flush=True)

    # 分析分镜
    print(f"\n[2/3] GPT分析分镜中...", flush=True)
    cfg = Config()
    analysis = analyze_storyboard(storyboard_text, "EP30_分镜", cfg)

    print(f"✅ 分析完成!", flush=True)
    print(f"  主基调: {analysis['primary_tone']}", flush=True)
    print(f"  总时长: {analysis['total_duration_s']}s", flush=True)
    print(f"  情绪分段数: {len(analysis.get('cues', []))}", flush=True)
    for cue in analysis.get('cues', []):
        print(f"    - {cue.get('start_s', 0)}-{cue.get('end_s', 0)}s: {cue.get('tone', '')}", flush=True)

    # 生成BGM
    print(f"\n[3/3] 开始生成多段BGM...", flush=True)
    print(f"注意: 每段生成之间会等待30秒以避免API速率限制", flush=True)

    result = generate_bgm_multisegment(
        analysis=analysis,
        title="EP30_分镜",
        out_dir=output_dir,
        config=cfg,
    )

    print(f"\n{'=' * 80}", flush=True)
    print(f"✅ 生成完成!", flush=True)
    print(f"  最终音频: {result['final_audio']}", flush=True)
    print(f"  生成段数: {result['unique_segments']}", flush=True)
    print(f"  情绪类型: {result['unique_tones']}", flush=True)
    print(f"{'=' * 80}", flush=True)

if __name__ == "__main__":
    main()
